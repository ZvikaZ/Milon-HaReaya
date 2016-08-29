# -*- coding: utf-8 -*-
'''
module responsible for parsinf the runs of a paragraph
'''

import sizes
import docx
import docx_fork_ludoo
from text_segments import MilonTextSegments as TS, fake
import text_segments as ts
from docx2abstract_doc import *
import fixes

class MilonDocxParser():
	def __init__(self, doc_file_name, debug_file_name):
		self.unknown_list = []
		self.temp_l = []
		self.debug_file_name = debug_file_name
		self.word_doc = docx.Document(doc_file_name)
		self.word_doc_footnotes = docx_fork_ludoo.Document(doc_file_name)

	def finish(self):
		if self.unknown_list:
			print "\n\nMissing:"
			print parse.unknown_list

	def paragraphs(self):
		'''
		generator of paragraph-footnotes tuples
		'''
		# TODO - size_kind - find a way to get rid of it.
		with open(self.debug_file_name, 'w') as debug_file:
			for (paragraph, footnote_paragraph) in zip(self.word_doc.paragraphs, self.word_doc_footnotes.paragraphs):
				para, footnotes, size_kind = self.parse_runs(paragraph, footnote_paragraph, debug_file)
				yield (para, footnotes, size_kind)

	def parse_runs(self, paragraph, footnote_paragraph, debug_file): # TODO debug_file - remove
		if not paragraph.text.strip():
			return [], [], None

		size_kind = None # when the return statment is changed, this line should be deleted
		footnotes = []
		para = []
		debug_file.write("\n\nNEW_PARA:\n------\n")
		for (run, footnote_run) in zip(paragraph.runs, footnote_paragraph.runs):
			s = "!%s.%s:%s$" % (run.style.style_id, docxCode2segType.get(run_style_id(run), run_style_id(run)), run.text)
			# print "!%s:%s$" % (docxCode2segType.get(run.style.style_id, run.style.style_id), run.text)
			debug_file.write(s.encode('utf8') + ' ')
			type = docxCode2segType.get(run_style_id(run), TS.unknown)

			if run.font.size and run.text.strip():
				size_kind = sizes.match(run.font.size)
				if size_kind == TS.unknown:
					print "!%s. Size: %d, Bool: %s, %s:%s$" % (size_kind, run.font.size, run.bold, type, run.text)
				if size_kind not in ('normal', TS.unknown):
					type = size_kind

			if 'unknown' in type and run.text.strip():
				type = fixes.fix_unknown(run)

			elif type == "DefaultParagraphFont":
				type = fixes.fix_DefaultParagraphFont(run)
				# print paragraph.style.style_id, run.bold, run.font.size, s

			elif run.bold:
				type = self.bold_type(s, type, run)

			# single run & alignment is CENTER and ...-> letter heading
			elif len(paragraph.runs) == 1 and paragraph.alignment is not None and int(paragraph.alignment) == 1\
		            and "heading" not in type and run.text.isalpha():
				# print "NEW heading letter!", s
				size_kind = "heading_letter"
				type = size_kind


			try:
				if run.element.rPr.szCs is not None and run.text.strip():
					type = fixes.fix_sz_cs(run, type, debug_file)

				if run.element.rPr.bCs is not None and run.text.strip():
					type = fixes.fix_b_cs(run, type)

				# NOTE: this footnote number need no fix.
				# it is a recurrance, therefore it has no id.
				if is_footnote_recurrence(run, type):
					type = TS.footnoteRec

			except:
				pass


			para.append((type, run.text))

			if type == TS.unknown:
				if run_style_id(run) not in unknown_list:
					unknown_list.append(run_style_id(run))
					print paragraph.text
					s = "\nMissing: !%s:%s$\n\n" % (run_style_id(run), run.text)
					print s
					debug_file.write(s.encode('utf8') + ' ')

			footnote_references = footnote_run.footnote_references
			if footnote_references:
				for note in footnote_references:
					para.append((TS.footnote, str(note.id)))
					footnotes.append(note.id)

		para.append((TS.newLine, '\n'))
		para = fixes.analyze_and_fix(para)
		return para, footnotes, size_kind # TODO this should be changed. size_kind shouldn't be returned and moved to other functions. it makes no sence.

	def bold_type(self, s, type, run):
		if type == TS.definitionNormal:
			return TS.subjectSmall
		elif type == TS.sourceNormal and run.style.style_id == "s03":
			return TS.subSubjectSmall
		elif type == TS.definitionSmall and run.style.style_id == "s05":
			return TS.subSubjectSmall
		elif type == TS.sourceNormal and run.style.style_id == "DefaultParagraphFont" and run.font.size == 139700:
			return TS.subjectNormal
		elif type == TS.sourceNormal and run.style.style_id == "DefaultParagraphFont" and run.font.size != 139700:
			return TS.subSubjectNormal
		elif type == TS.unknownLight and run.style.style_id == "s04" and run.font.size == 114300:
			return TS.subjectLight
		elif type == TS.unknownLight and run.style.style_id == "s04" and run.font.size == 101600:
			return TS.subSubjectLight
		elif type == TS.definitionLight and run.style.style_id == "s12" and run.font.size == 101600:
			return TS.subSubjectLight
		elif type == TS.definitionLight and run.style.style_id == "s12" and run.font.size is None:
			# TODO - verify that it's always OK
			return TS.subSubjectLight
		elif type == TS.sourceNormal:
			print "Strange TS.sourceNormal bold!"
		elif 'subject' in type or 'heading' in type:
			return type
		elif run.text.strip() in (u"◊", "-", ""):
			return type
		else:
			if type not in self.temp_l:
				print "Unexpected bold!", type
				print s, type, run.text, run.font.size
				assert False
				self.temp_l.append(type)
			return type


##########################################################################################
## Static methods ########################################################################
##########################################################################################

# support old and new version of docx
# new is preferred, because, well, it's new...
# old is preferred because I'm using a branch with footnotes support
def run_style_id(run):
    try:
        return run.style.style_id
    except:
        if run.style:
            return run.style
        else:
            return 'DefaultParagraphFont'

def is_footnote_recurrence(run, type):
    # a number in superscript, that's not defined as a footnote
    return \
        run.element.rPr.vertAlign is not None \
        and type != TS.footnote \
        and run.text.strip().isdigit() \
        and run.element.rPr.vertAlign.values()[0] == 'superscript'
