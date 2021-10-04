# -*- coding: utf-8 -*-
""" Currently, this is the 'main' module of this project.
By default, it parses 'dict.docx' (that contains Milon HaReaya's source Word file,
creates 'html_docs_l' - internal representation of all the Milon,
'subjects_db' which is used for searching (and is written as JSON file for the JS)
and then creates 'output/' with a working HTML/CSS/JS site, and zips it to 'milon.zip'

If then creates an apk file and Electron .exe NSIS installer,
and pushes the APK (automatically) to Google Play.
"""

# TODO: refactor, split to files, unit tests
# TODO: why "Pashut" is unknown?
# TODO: fix "FOOTNOTE undefined: af7 None  :  homo"

# TODO: TEX: with new version, " looks different (it's curly)

# TODO: TEX: 'tex.full' line 8089 failure (8246) - because of "Vav" in "Otiyot" - need to delete previous line, and re-enter
# TODO: TEX: Mehkarim UVeurim - handle style (w/o numbers...)
# TODO: TEX: add prefixes and appendices
# TODO: TEX: clean milon.tex, handle koma recommendations
# TODO: TEX: make sure that "fake_subject_normal" is correct - currently different between HTML and LyX
# TODO: TEX: make sure there aren't any "Takala" in .tex file
# TODO: TEX: publish my Tex packages?
# TODO: create indexes? (both before and after main text...)

# TODO: Clean 'UNKNOWN's and 'fix_sz_cs'
# TODO: verify that it's running on clean GIT clone

# TODO: change 'is_prev_subject(..)' to correctly handle "Toar Shem Tov" - should be more freely checking
# TODO: otiyot - stam font
# TODO: pagination at end
# TODO: subject_light vs sub-subjet_light - wait for Rav's response

# TODO: subjects size in Mehkarim
# TODO: references numbering
# TODO: search "Natziv" not working
# TODO: Yud and Lamed in Psukim
# TODO: splitted bubject, like "אמר לו הקדוש ברוך הוא (לגבריאל° שבקש להציל את אברהם־אבינו° מכבשן האש) אני יחיד בעולמי והוא יחיד בעולמו, נאה ליחיד להציל את היחיד"

# TODO: circles shouldn't be part of subjects (and what about parentheses?)
# TODO: "Mehkarim" - make links, check styles!
# TODO: Change "opening_abbrev.html" styling
# TODO: handle footnotes' styles
# TODO: "Ayen", "Re'e" - see mail from 22.1.16
# TODO: "all subjects" page

# TODO: make headings to links
# TODO: save current location (and history?, with back and forward?) - use HTML5 local storage
# TODO: search - results page - write first words of definition
# TODO: add letters to TOC
# TODO: make smarter links on circles ('Oneg' with and w/o Vav, 'zohama' with Alef or He, etc.)
# TODO: increase/decrease font size
# TODO: make definition in new line? (without ' - ')

# TODO: replace menu with Bootstrap style menu
# TODO: Make index.html's links clickable, or copyable
# TODO: better icon
# TODO: iphone?
# TODO: GUI

# reminder:
# dropped PhonaGap for Monaca:
# https://console.monaca.mobi/dashboard
# (maybe change Rimon...)


# remember:
# http://stackoverflow.com/questions/10752055/cross-origin-requests-are-only-supported-for-http-error-when-loading-a-local

# hopefully, I will get a new delivery of python-docx supporting szCs
# (see https://github.com/python-openxml/python-docx/issues/248 )
# in the meanwhile, I've hacked it locally
import sys

import build_cordova
import build_electron
import search_index
from misc import replace_in_file
from search_index import calc_subject_id, clean_name

sys.path.insert(0, r'C:\Zvika\PycharmProjects\python-docx')
sys.path.insert(0, r'C:\Users\sdaudi\Github\python-docx')
import docx     # from aforementioned path

import docx_fork_ludoo
import dominate
import dominate.tags as tags
import re
import zipfile
import os
import shutil
import html.parser
import json
import copy

import upload_google_play
import htmler
import texer

html_parser = html.parser.HTMLParser()

# process = "Full"
# process = "Compile"
process = "ZIP"

if process in ["Full", "Compile"]:
    doc_file_name = 'מילון הראיה.docx'
    create_html = True
    #create_latex = True
    create_latex = False
else:
    # doc_file_name = 'dict_few.docx'
    doc_file_name = 'dict_check.docx'
    # doc_file_name = 'dict_short.docx'
    # doc_file_name = 'dict_half.docx'
    # doc_file_name = 'מילון הראיה.docx'

    create_html = True
    create_latex = False

    # create_html = False
    # create_latex = True

word_doc = docx.Document(os.path.join('input_dict', doc_file_name))
word_doc_footnotes = docx_fork_ludoo.Document(os.path.join('input_dict', doc_file_name))


# support old and new version of docx
# new is preferred, because, well, it's new...
# old is preferred because I'm using a branch with footnotes support
def run_style_id(run):
    bold = ""
    if run.text.strip() and run.font.cs_bold:
        # print "BOLD: ", run.bold, run.font.bold, run.font.cs_bold, ": ", run.text
        bold = "_bold"
    try:
        return run.style.style_id + bold
    except:
        if run.style:
            return run.style + bold
        else:
            return 'DefaultParagraphFont' + bold


styles = {
    's01': 'subject_normal',
    's11': 'sub-subject_normal',
    's03_bold': 'sub-subject_small',    #?Fixing appendix  ?
    's02': 'definition_normal',
    's02_bold': 'definition_normal',     #?Fixing appendix
    's03': 'definition_small',
    'Heading3Char': 'definition_normal',
    '1': 'definition_normal',   #?
    'FootnoteTextChar1': 'definition_normal',   #?
    'HebrewChar': 'definition_normal',   #?

    # this is problematic! has its own function to handle it
    'DefaultParagraphFont': 'DefaultParagraphFont',
    'DefaultParagraphFont_bold': 'DefaultParagraphFont',    #TODO: correct?

    's15': 'subject_small',
    's17': 'subject_small',
    's1510': 'subject_small',
    's1510_bold': 'subject_small',
    's05': 'definition_small',
    's05_bold': 'sub-subject_small',    #?Fixing appendix
    's038': 'source_small',
    's0590': 'source_small',
    's050': 'source_small',
    '050': 'source_small',

    's149': 'subject_light',
    's14': 'subject_light',
    's16': 'sub-subject_light',
    's12_bold': 'sub-subject_light',
    's168': 'sub-subject_light',
    's048': 'definition_light',
    's12': 'definition_light',
    's04': 'unknown_light',
    's127': 'source_light',

    's02Symbol': 's02Symbol',   # MeUyan

    'FootnoteReference': 'FootnoteReference',
    'EndnoteReference': 'EndnoteReference', #?

    # 17.3.19 - these were added with new Office 365 (2019) - seem like new aliases for existing styles
    'a0': 'DefaultParagraphFont',
    'a0_bold': 'DefaultParagraphFont',
    'a3': 'FootnoteReference',
    'a7': 'EndnoteReference',
    '30': 'definition_normal',
    '11': 'definition_normal',
    'a6': 'definition_normal',
    # \17.3.19
}


# if the actual size is greater
class Sizes:
    my_dict = {
        381000: 'heading_title',
        330200: 'heading_section',              # e.g., "Tora"
        279400: 'heading_sub-section-bigger',   # e.g., "Mehkarim Beurim"
        215900: 'heading_sub-section',          # e.g., "Avraham Yitzhak VeYaakov"
        177800: 'heading_letter',
        165100: 'normal',               # 152400
    }

    # yeah, it's not nice and programmaticish to have this twice
    # but it's more efficient :)
    normal = 165100

    def match(self, size):
        if size > self.normal:
            return self.my_dict.get(size, "unknown")
        else:
            return 'normal'

    def get_heading_type(self, kind):
        assert 'heading' in kind
        if kind == 'heading_title':
            return tags.h1
        elif kind == 'heading_section':
            return tags.h2
        elif kind == 'heading_sub-section-bigger' or kind == 'section_title_secondary':
            return tags.h3
        elif kind == 'heading_sub-section':
            return tags.h4
        elif kind == 'heading_letter':
            return tags.h5
        else:
            assert False

sizes = Sizes()

unknown_list = []

#TODO remove this
# dictionary mapping subjects to list of pointers
# each pointer is a tuple of (subject, html_doc's section name, url)
subjects_db = {}


def subject(html_doc, type, text):
    result = tags.span()
    with result:
        clean_text = clean_name(text.strip())
        new_subject_l = subjects_db.get(clean_text, [])
        subject_id = calc_subject_id(text.strip(), len(new_subject_l))
        new_subject_l.append((text.strip(), html_doc.section, "%s.html#%s" % (html_doc.index, subject_id)))
        subjects_db[clean_text] = new_subject_l

        with tags.span(text, id=subject_id):
            tags.attr(cls=type)

        # with tags.span(tags.a(text, href="#%s" % text.strip(), id=text.strip())):
        #     tags.attr(cls=type)
    return result


def regular(type, text):
    result = tags.span()
    with result:
        if type in ['footnote', 'footnote_recurrence']:
            with tags.a("(%s)" % text.strip()):
                tags.attr(cls="ptr")
        else:
            if "\n" in text:
                print("New:", text)
            if "°" in text:
                href = re.sub("°", "", text)
                href = re.sub("־", " ", href)
                with tags.span(tags.a(text, href="#"+href)):
                    tags.attr(cls=type)
            else:
                with tags.span(text):
                    tags.attr(cls=type)
    return result

def is_footnote_recurrence(run, type):
    # a number in superscript, that's not defined as a footnote
    return \
        run.element.rPr.vertAlign is not None \
        and type != 'footnote' \
        and run.text.strip().isdigit() \
        and list(run.element.rPr.vertAlign.values())[0] == 'superscript'

def is_subject(para, i, next=False):
    type, text = para[i]
    # print "is? ", type, text.strip()
    # if 'subject' in type and not re.search(r"\w", text, re.UNICODE):
    #     print "!", text
    # print "is?", type, text, i, ('subject' in type and re.search(r"\w", text, re.UNICODE))

    # if 'subject' in type and re.search(r"\w", text, re.UNICODE) and i>0:
    #     p_type, p_text = para[i-1]
    #     print "?", i-1, p_type, p_text
    return 'subject' in type and re.search(r"\w", text, re.UNICODE) and 'fake' not in type

def is_prev_subject(para, i):
    try:
        return (is_subject(para, i-2) and
                (para[i-1][1].replace('"','').strip() == "-") or (para[i-1][0] == "footnote"))
    except:
        return False

def is_prev_newline(para, i):
    try:
        return para[i-1][0] == "new_line" or (para[i-2][0] == "new_line" and para[i-1][1] == "")
    except:
        return False

def is_prev_meuyan(para, i):
    try:
        return para[i-1][0] == "s02Symbol"
    except:
        return False



def make_sub_subject(subj):
    if subj == 'subject_small':
        return 'sub-subject_normal'
    else:
        return subj

def is_subject_small_or_sub_subject(s):
    return s in ['subject_small', 'sub-subject_normal']

def analyze_and_fix(para):
    # unite splitted adjacent similar types
    prev_type, prev_text = None, ""
    new_para = []
    for (raw_type, text_raw) in para:
        text = text_raw.replace("@", "")
        if text == "◊":
            if prev_text == "":
                # new paragraph with meuayn
                type = "centered_meuyan"
                # see below, end of function, we're going to fix it...
            else:
                # regular meuyan, doesn't need to be centerized
                type = "s02Symbol"
        else:
            type = raw_type
        if prev_type:
            if (type == prev_type) or \
                    (is_subject_small_or_sub_subject(type) and is_subject_small_or_sub_subject(prev_type)) or \
                    (prev_type != "footnote" and prev_type != "footnote_recurrence" and text.strip() in ("", "°", "־", ",")):
                prev_text += text
            elif prev_type == "centered_meuyan" and type == "s02Symbol":
                prev_text += text
            else:
                new_para.append((prev_type, prev_text))
                prev_type, prev_text = type, text
        else:
            prev_type, prev_text = type, text
    new_para.append((prev_type, prev_text))

    # make new_lines stand on their own
    para = new_para
    new_para = []
    for (type, text) in para:
        lines = text.split("\n")
        if len(lines) > 1:
            for (i, line) in enumerate(lines):
                if line:
                    new_para.append((type, line))
                if i+1 < len(lines):
                    new_para.append(("new_line", "\n"))
        else:
            new_para.append((type, text))

    # fix wrong subjects
    para = new_para
    new_para = []
    for (index, (type, text)) in enumerate(para):
        # if 'subject' in type:
        if is_subject(para, index):
            # real subject is either:
            # first
            # after new_line and empty
            # after subject,"-"
            # after Meuyan
            if (index == 0) or (is_prev_newline(para, index)) or (is_prev_meuyan(para, index)):
                new_para.append((type, text))
            elif (is_prev_subject(para, index)):
                new_para.append((make_sub_subject(type), text))
            elif new_para[index-1][0] in ('sub-subject_normal', 'subject_small'):
                new_para.append((make_sub_subject(type), text))
            else:
                new_para.append(("fake_"+type, text))
        elif 'subject' in type:
            # it's got a subject, but 'is_subject' failed
            new_para.append(("fake_"+type, text))
        else:
            new_para.append((type, text))

    # fix missing 'source's
    para = new_para
    new_para = []
    source_pattern = re.compile(r"(.*)(\[.*\])(\..*)")
    for (type, text) in para:
        if source_pattern.match(text) and not 'source' in type:
            g = source_pattern.match(text)
            new_para.append((type, g.group(1)))
            if type in ['definition_small', 'fake_sub-subject_small']:
                new_para.append(('source_normal', g.group(2)))
            elif type == 'definition_normal':
                new_para.append(('source_small', g.group(2)))
            elif 'light' in type:
                new_para.append(('source_light', g.group(2)))
            else:
                print("Fix missing 'source's, unknown type: ", end='')
                print(type, g.group(2))
                new_para.append((type, g.group(2)))
            new_para.append((type, g.group(3)))
        else:
            new_para.append((type, text))

    # make links from circles - °
    para = new_para
    new_para = []
    pattern = re.compile("([\S־]*\S+°)", re.UNICODE)
    for (type, text) in para:
        # don't do this for subjects - it complicates their own link name...
        if "°" in text and 'subject' not in type:
            for (chunk) in pattern.split(text):
                new_para.append((type, chunk))
        else:
            new_para.append((type, text))


    # fix new lines inside headings
    para = new_para
    new_para = []
    ignore_new_line = False
    for (type, text) in para:
        if 'heading' in type:
            ignore_new_line = True
            new_para.append((type, text))
        elif type == "new_line" and ignore_new_line:
            pass
        else:
            new_para.append((type, text))


    # scan for 'empty subjects' ...
    has_subject = False
    has_definition = False
    for (type, text) in para:
        if 'subject' in type and 'fake' not in type and text.strip():
            has_subject = True
            has_definition = False
        elif 'definition' in type:
            has_definition = True

    # ... and fix 'em if required
    if has_subject and not has_definition:
        # empty subject
        #TODO - might be caused by 'heading' interpreted as subject
        para = new_para
        new_para = []
        for (type, text) in para:
            if 'subject' in type and 'fake' not in type:
                new_para.append(('fake_'+type, text))
            else:
                new_para.append((type, text))


    # 'centered_meuyan' is legal only if it's the only "heavy" thing in the paragraph, otherwise, it's a regular 'meuyan"
    found_centered_meuyan = False
    should_replace_centered_meuayn = False
    for (type, text) in para:
        if type == "centered_meuyan":
            found_centered_meuyan = True
        elif type != "new_line" and found_centered_meuyan:
            should_replace_centered_meuayn = True
    if found_centered_meuyan:
        para = new_para
        new_para = []
        if should_replace_centered_meuayn:
            for (type, text) in para:
                if type == "centered_meuyan":
                    new_para.append(("s02Symbol", text))
                else:
                    new_para.append((type, text))
        else:
            # it's a centered meuyan, and should stay such; but we don't need anything else in it
            for (type, text) in para:
                if type == "centered_meuyan":
                    new_para.append((type, text))
                elif type == "new_line":
                    # redundant, and causes troubles...
                    pass
                else:
                    print("Unexpected type in centered meuyan")
                    print(type, text)
                    assert False



    with open('output/debug_fix.txt', 'a', encoding='utf-8') as debug_file:
        debug_file.write("---------------\n")
        for (type, text) in new_para:
            s = "%s:%s.\n" % (type, text)
            debug_file.write(s + ' ')

    # fix
    return new_para

# return True if updated
def update_values_for_href(child, href):
    values = subjects_db.get(href)
    #TODO: support showing more than 1 result
    if values is None:
        return False
    if len(values) == 1:
        _, _, url = values[0]
        child.children[0]['href'] = url
        return True
    elif len(values) > 1:
        # # tuple of (subject, html_doc's section name, url)
        # subject, section, url = values[]
        child.children[0]['href'] = "search.html?method=exact_subject&term="+href
        return True
    else:
        assert False

def update_href_no_link(child):
    assert len(child.children[0].children) == 1
    text = html_parser.unescape(child.children[0].children[0])
    child.children[0] = tags.span(text)


def fix_links(html_docs_l):
    # fix outbound links
    print("Fixing links")
    for (doc) in html_docs_l:
        for (parent) in doc.body.children[0].children:
            for middle in parent.children:
                try:
                    for child in middle.children:
                        if 'definition' in child.attributes.get('class', ()):
                            href = ""
                            try:
                                href = child.children[0].attributes.get("href")
                            except AttributeError as e:
                                pass

                            # it's a link - try to update it
                            if href:
                                # first, strip it of weird chars
                                try:
                                    href = clean_name(href)

                                    updated = False
                                    if update_values_for_href(child, href):
                                        updated = True
                                    else:
                                        if href[0] in ("ה", "ו", "ש", "ב", "כ", "ל", "מ"):
                                            updated = update_values_for_href(child, href[1:])
                                    if not updated:
                                        # failed to update - it's not a real link...
                                        update_href_no_link(child)

                                except Exception as e:
                                    pass
                                    print(e, "Exception of HREF update", href)
                                    #TODO - investigate why it happens? (it's a single corner case, I think)
                except AttributeError:
                    pass

    def sorter(html_doc):
        if html_doc.name in ["ערכים כלליים",]:
            return "FIRST"
        if html_doc.name in ["נספחות",]:
            return "תתתתתתתתתת"    #last
        else:
            return html_doc.name

    # update sections menu
    for (doc) in html_docs_l:
        letters_l = []

        content_menu = doc.body.children[0].children[1].children[0].children[-1].children[-1]
        assert content_menu['class'] == 'dropdown-menu dropdown-menu-left scrollable-menu'

        with content_menu:
            with tags.li():
                tags.a("אודות", href="index.html")
            with tags.li():
                tags.a("הקדמות", href="opening_intros.html")
            with tags.li():
                tags.a("הסכמות", href="opening_haskamot.html")
            with tags.li():
                tags.a("קיצורים", href="opening_abbrev.html")
            with tags.li():
                tags.a("סימנים", href="opening_signs.html")
            # if you add more entries here, please update add_menu_to_apriory_htmls(..)
            with tags.li():
                tags.attr(cls="divider")

            sorted_html_docs_l = sorted(html_docs_l, key=sorter)

            for (html_doc) in sorted_html_docs_l:
                # Only if this a 'high' heading, and not just a letter - include it in the TOC
                if html_doc.name != "NEW_LETTER":
                    with tags.li():
                        tags.a(html_doc.name, href=str(html_doc.index)+".html")
                        if doc.section == html_doc.section:
                            tags.attr(cls="active")
                else:
                    # it's a letter - if it's related to me, save it
                    if doc.section == html_doc.section:
                        letters_l.append(html_doc)

        with doc.body.children[-1]:
            assert doc.body.children[-1]['class'] == 'container-fluid'
            with tags.ul():
                tags.attr(cls="pagination")
                for (html_doc) in letters_l:
                    tags.li(tags.a(html_doc.letter, href=str(html_doc.index)+".html"))



    return html_docs_l


new_lines_in_raw = 0


def add_running_tag(body, running_tag):
    if str(running_tag) != "<span></span>":
        search_index.learn(running_tag, html_doc)
        body += running_tag
    return (body, tags.span())


def add_to_output(html_doc, para):
    global new_lines_in_raw
    # we shouldn't accept empty paragraph (?)
    assert len(para) > 0

    body = html_doc.body.children[-1]
    assert body['class'] == 'container-fluid';
    running_tag = tags.span()
    for (i, (type, text)) in enumerate(para):
        if 'heading' in type and text.strip():
            # tags.p()
            # tags.p()
            body, running_tag = add_running_tag(body, running_tag)
            heading = sizes.get_heading_type(type)
            print(type, text)
            body += heading(text)
        elif type == "new_line":
            new_lines_in_raw += 1
            if new_lines_in_raw == 1:
                running_tag += tags.br()
            elif new_lines_in_raw == 2:
                body, running_tag = add_running_tag(body, running_tag)
                body += tags.p()
            else:
                pass
        elif is_subject(para, i):
            if not is_prev_subject(para, i):
                # tags.p()
                #tags.br()
                pass

            body, running_tag = add_running_tag(body, running_tag)
            running_tag = subject(html_doc, type, text)
        else:
            running_tag += regular(type, text)

        if type != "new_line":
            new_lines_in_raw = 0

    # tags.br()
    body, running_tag = add_running_tag(body, running_tag)



def fix_sz_cs(run, type):
    result = type
    szCs = list(run.element.rPr.szCs.attrib.values())[0]

    try:
        eastAsia = 'eastAsia' in list(run.element.rPr.rFonts.attrib.values())
    except:
        eastAsia = False

    try:
        hint_cs = 'cs' in list(run.element.rPr.rFonts.attrib.values())
    except:
        hint_cs = False


    if szCs == "20" and 'subject' in type:
        if run.style.style_id == "s01":
            # s = "!Fixed!szCs=%s:%s!bCs=%s!" % (szCs, run.text, run.element.rPr.bCs.attrib.values()[0])
            s = "!Fixed!szCs=%s:%s!" % (szCs, run.text)
            # print s
            debug_file.write(s + ' ')
            # return 'definition_normal'
            return 'subject_small'
    elif szCs == "22" and type == 'definition_normal':
        return 'subject_normal'
    elif szCs == "22" and type == 'sub-subject_normal':
        return 'subject_normal'
    elif szCs == "14" and type == 'definition_light':
        return 'source_light'
    elif szCs == "16" and type == 'definition_small' and run.text.replace(',','').replace(':','').strip().isdigit():
        return 'definition_small'		# keep original type
    elif szCs == "16" and type == 'definition_small': # and hint_cs:
        return 'source_small'
    elif szCs == "14" and type == 'definition_small' and hint_cs:   # or maybe my 'isdigit'
        return 'source_small'
    elif szCs == "16" and type == 'source_normal' and eastAsia:
        return 'source_small'
    elif szCs == "16" and type == 'source_normal' and hint_cs:
        return 'source_small'
    elif szCs == "14" and type == 'source_normal' and hint_cs:
        return 'source_small'
    elif szCs == "18" and type == 'definition_normal':
        return 'definition_small'
    elif szCs == "18" and type == 'sub-subject_normal':
        # return 'sub-subject_normal'             #20.11.16 - Trying to fix Appendix' bold and fonts
        return 'sub-subject_small'
    elif szCs == "18" and type == 'subject_small':
        return 'sub-subject_small'
    elif szCs == "18" and type == 'subject_normal':
        return 'definition_small'
    elif szCs == "18" and type == 'source_normal':
        return 'definition_small'
    elif szCs == "20" and type == 'unknown_light':
        return 'definition_normal'
    elif szCs == "26" and type == 'subject_normal':
        ## wild guess, might break everything :-(
        ## be careful here...
        print("ZZ: Fixed to 'section_title_secondary': ", run.text.strip())
        return 'section_title_secondary'
    elif run.text.strip():
        #print("fix_sz_cs::Unsupported value: ", szCs, "type:", type, ". At: ", run.text)    #TODO: clean this!!!
        pass
    else:
        pass
    return result

def fix_b_cs(run, type):
    result = type
    try:
        bCs = list(run.element.rPr.bCs.attrib.values())[0]
        try:
            hint_cs = run.element.rPr.rFonts.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hint', None) == 'cs'
        except:
            hint_cs = False

        try:
            szCs = list(run.element.rPr.szCs.attrib.values())[0]
        except:
            szCs = None

        try:
            fonts = run.element.rPr.rFonts.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs', None)
        except:
            fonts = None

        if bCs == "0" and 'subject' in type:
            if (run.style.style_id == "s01" and fonts != "Narkisim" and (run.bold != True or hint_cs or szCs == '20')) or \
                  (run.style.style_id == "s11" and (run.bold != True or (bCs == '0' and fonts != "Narkisim"))):
                if type in ('subject_small', 'sub-subject_normal'):
                    return 'definition_normal'
            else:
                #print("Unknown b_cs=0")
                pass
    except:
        pass
    return result

def fix_misc_attrib(run, type):
    try:
        eastAsia = list(run.element.rPr.rFonts.attrib.values())[0] == 'eastAsia'
        if eastAsia:
            if type == "source_normal":
                return "definition_small"
            else:
                # print type, run.text
                return type
        else:
            return type


    except:
        return type



def fix_unknown(run):
    if run.font.size == 114300 and run.style.style_id == 's04':
        return 'subject_light'
    elif run.font.size == 101600 and run.style.style_id == 's04' and run.font.cs_bold:
        return 'sub-subject_light'
    elif run.font.size == 101600 and run.style.style_id == 's04' and not run.font.cs_bold:
        return 'definition_light'
    elif run.font.size is None and run.style.style_id == 's04':
        return 'definition_light'
    elif run.font.size == 88900 and run.style.style_id == 's04':
        return 'source_light'
    else:
        print("UNKNOWN: ", run.text)
        return 'unknown_light'


def fix_DefaultParagraphFont(run):
    # only if it's really a text
    if run.text.strip():
        if run.font.size == 330200:
            return 'heading_section'
        elif run.font.size == 177800 and run.font.cs_bold:
            return 'subject_normal'
        elif run.font.size == 152400 and run.font.cs_bold:
            return 'sub-subject_normal'
        elif run.font.size == 152400 and not run.font.cs_bold:
            return 'definition_normal'
        elif run.font.size == 139700 and run.font.cs_bold:
            return 'subject_normal'
        elif run.font.size == 139700 and not run.font.cs_bold:
            return 'definition_normal'
        elif run.font.size == 127000:
            return 'definition_normal'
        elif run.font.size == 114300 and run.font.cs_bold:
            return 'sub-subject_normal'
        elif run.font.size == 114300 and not run.font.cs_bold:
            return 'source_normal'
        elif run.font.size == 101600:
            return 'source_small'
        elif run.font.size == 88900:
            return 'source_small'
        elif run.font.size is None and run.font.cs_bold:
            return 'sub-subject_normal'
        elif run.font.size is None and not run.font.cs_bold:
            return 'definition_normal'
        else:
            if run.text.strip() not in ("-", "(", ")", "[", "]", "'", '"', ","):
                print("AH!", ":",run.text.strip(),".", run.font.size, run.bold, run.font.cs_bold)
                assert False
            #else:
            #    return 'DefaultParagraphFont'
    else:
        return 'DefaultParagraphFont'

temp_l = []


def uniqify(s):
    return "".join(set(s))


def bold_type(s, type, run):
    if type == 'definition_normal':
        return 'subject_small'
    elif type == 'source_normal' and run.style.style_id == "s03":
        return 'sub-subject_small'
    elif type == "definition_small" and run.style.style_id == "s05":
        return 'sub-subject_small'
    elif type == 'source_normal' and run.style.style_id == "DefaultParagraphFont" and run.font.size == 139700:
        return 'subject_normal'
    elif type == 'source_normal' and run.style.style_id == "DefaultParagraphFont" and run.font.size != 139700:
        return 'sub-subject_normal'
    elif type == 'unknown_light' and run.style.style_id == "s04" and run.font.size == 114300:
        return 'subject_light'
    elif type == 'unknown_light' and run.style.style_id == "s04" and run.font.size == 101600:
        return 'sub-subject_light'
    elif type == 'definition_light' and run.style.style_id == "s12" and run.font.size == 101600:
        return 'sub-subject_light'
    elif type == 'definition_light' and run.style.style_id == "s04" and run.font.size is None:
        return 'subject-light'
    elif type == 'definition_light' and run.style.style_id == "s12" and run.font.size is None:
        # TODO - verify that it's always OK
        return 'sub-subject_light'
    elif type == 'source_normal':
        print("Strange 'source_normal' bold!")
    elif 'subject' in type or 'heading' in type:
        return type
    elif uniqify(run.text.strip()) in ("◊", "-", ""):
        return type
    else:
        if type not in temp_l:
            print("Unexpected bold!", type)
            print(s, type, run.text, run.font.size)
            assert False
            temp_l.append(type)
        return type


def open_html_doc(name, letter=None):
    html_doc = dominate.document(title="מילון הראיה")
    html_doc['dir'] = 'rtl'
    with html_doc.head:
        with tags.meta():
            tags.attr(charset="utf-8")
        with tags.meta():
            tags.attr(name="viewport", content="width=device-width, initial-scale=1")

        tags.script(src="jquery/dist/jquery.min.js")
        tags.link(rel='stylesheet', href='bootstrap-3.3.6-dist/css/bootstrap.min.css')
        tags.link(rel='stylesheet', href='style.css')
        tags.script(src="bootstrap-3.3.6-dist/js/bootstrap.min.js")
        tags.link(rel='stylesheet', href="bootstrap-rtl-3.3.4/dist/css/bootstrap-rtl.css")
        tags.link(rel='stylesheet', href='html_demos-gh-pages/footnotes.css')
        tags.script(src="milon.js")
        tags.script(src="html_demos-gh-pages/footnotes.js")
        tags.script(src="subjects_db.json")
        tags.script(src="index.json")
        tags.script(src="elasticlunr.min.js")




    html_doc.footnote_ids_of_this_html_doc = []
    html_doc.name = name
    if letter:
        html_doc.letter = letter
        html_doc.section = html_docs_l[-1].section
    else:
        html_doc.section = name

    html_doc.index = len(html_docs_l) + 1
    with html_doc.body:
        with tags.div():
            tags.attr(cls="container-fluid")
            # TODO: call page_loaded to update saved URL also in other links
            tags.script("page_loaded('%s.html')" % html_doc.index)

            with tags.div():
                tags.attr(cls="fixed_top_left", id="menu_bar")
                with tags.div():
                    with tags.button('הקודם', type="button"):
                        tags.attr(id="back_icon_button", type="button", cls="btn btn-default", onclick="goBack()")
                        # with tags.span():
                        #     tags.attr(cls="glyphicon glyphicon-arrow-left")
                    with tags.button('הבא', type="button"):
                        tags.attr(id="forward_icon_button", type="button", cls="btn btn-default", onclick="goForward()")
                        # with tags.span():
                        #     tags.attr(cls="glyphicon glyphicon-arrow-right")
                    with tags.button(type="button"):
                        tags.attr(id="search_icon_button", type="button", cls="btn btn-default")
                        with tags.span():
                            tags.attr(cls="glyphicon glyphicon-search")
                    with tags.span():
                        tags.attr(cls="dropdown")
                        with tags.button(type="button", cls="btn btn-primary") as b:
                            tags.attr(href="#") #, cls="dropdown-toggle")
                            with tags.span():
                                tags.attr(cls="glyphicon glyphicon-menu-hamburger")
                                # b['data-toggle'] = "dropdown"
                        with tags.ul():
                            tags.attr(cls="dropdown-menu dropdown-menu-left scrollable-menu")





    return html_doc


def close_html_doc(html_doc):
    with html_doc.body.children[-1]:
        assert html_doc.body.children[-1]['class'] == 'container-fluid'
        with tags.div(id="search_modal"):
            tags.attr(cls="modal fade")


    with html_doc:
        # add footnotes content of this section:
        with tags.ol(id="footnotes"):
            for (id) in html_doc.footnote_ids_of_this_html_doc:
                footnote = word_doc_footnotes.footnotes_part.notes[id + 1]
                assert footnote.id == id
                htmler.add_footnote_to_output(footnote.paragraphs)

        # add placeholder for searching
        tags.comment("search_placeholder")

    place_holder = "<!--search_placeholder-->"
    with open("input_web/stub_search.html", 'r', encoding='utf-8') as file:
        search_html = file.read()

    html_doc_name = html_doc.index
    name = "debug_%s.html" % html_doc_name
    with open("output/www/" + name, 'w', encoding='utf-8') as f:
        f.write(html_doc.render(pretty=True))
    replace_in_file("output/www/" + name, place_holder, search_html)

    name = "%s.html" % html_doc_name
    with open("output/www/" + name, 'w', encoding='utf-8') as f:
        f.write(html_doc.render(pretty=False))
        print("Created ", name)
    replace_in_file("output/www/" + name, place_holder, search_html)


heading_back_to_back = False
pattern = re.compile(r"\W", re.UNICODE)

# returns:
#  None - if no need for new section
#  string - with name of new required section
#  ('UPDATE_NAME', string) - to replace the name of newly opend section
#  ('NEW_LETTER', string) - if needs a new section, but w/o putting it in the main TOC
def is_need_new_section(para, prev_name):
    global heading_back_to_back
    for (type, text) in para:
        if type in ("heading_title", "heading_section"):
            if not heading_back_to_back:
                heading_back_to_back = True
                return text.strip()
            else:
                # the previous, and this, are headings - unite them
                if prev_name != "מדורים":
                    result = prev_name + " " + text.strip()
                else:
                    # in the special case of 'Section' heading - we don't need it
                    result = text.strip()
                return ('UPDATE_NAME', result)
        elif type == "heading_letter":
            return ('NEW_LETTER', text.strip())

    # if we're here - we didn't 'return text' with a heading
    heading_back_to_back = False

def fix_section_name(name):
    if name == "מילון הראיה":
        return "ערכים כלליים"
    else:
        return name

html_docs_l = []
def get_active_html_doc(para):
    try:
        prev_name = html_docs_l[-1].name
    except:
        prev_name = None
    name = is_need_new_section(para, prev_name)
    if name:
        if isinstance(name, tuple):
            op, new = name
            if op == 'NEW_LETTER':
                html_docs_l.append(open_html_doc(op, letter=new))
            else:
                assert op == 'UPDATE_NAME'
                print("Updating ", new)
                html_docs_l[-1].name = new
                html_docs_l[-1].section = new
        else:
            fixed_name = fix_section_name(name)
            html_docs_l.append(open_html_doc(fixed_name))
    return html_docs_l[-1]



try:
    shutil.rmtree("output")
except FileNotFoundError:
    pass

try:
    shutil.rmtree("tex")
except FileNotFoundError:
    pass

os.mkdir("output")
os.mkdir("output/www")
os.mkdir("output/www/html_demos-gh-pages")
os.mkdir("tex")

os.chdir("input_web")
for (f) in (
    'style.css',
    'html_demos-gh-pages/footnotes.css',
    'html_demos-gh-pages/footnotes.js',
    'milon.js',
    'elasticlunr.min.js',
    'index.html',
    'opening_abbrev.html',
    'opening_haskamot.html',
    'opening_intros.html',
    'opening_signs.html',
    'search.html',
):
    shutil.copyfile(f, os.path.join("../output/www/", f))

shutil.copytree('res', '../output/res')

for (d) in (
    'bootstrap-3.3.6-dist',
    'bootstrap-rtl-3.3.4',
    'jquery',
):
    shutil.copytree(d, os.path.join("../output/www", d))

os.chdir("../input_tex")
for (f) in (
    "milon.tex",
    "polythumbs.sty",
    "hebcolumnbal.sty",
    #  "hebrew-gymatria-fix.sty",     # Rav Kalner asked not to do it. Leaving it here for future reference...
):
    shutil.copyfile(f, os.path.join("../tex", f))
os.chdir("../")


texer.open_latex()
# Here starts the action!
with open('output/debug.txt', 'w', encoding='utf-8') as debug_file:
    for (paragraph, footnote_paragraph) in zip(word_doc.paragraphs, word_doc_footnotes.paragraphs):
        if paragraph.text.strip():
            # print "Paragraph:", paragraph.text, "$"
            para = []
            debug_file.write("\n\nNEW_PARA:\n------\n")
            for (run, footnote_run) in zip(paragraph.runs, footnote_paragraph.runs):
                s = "!%s.%s:%s$" % (run.style.style_id, styles.get(run_style_id(run), run_style_id(run)), run.text)
                # print "!%s:%s$" % (styles.get(run.style.style_id, run.style.style_id), run.text)
                debug_file.write(s + ' ')
                type = styles.get(run_style_id(run), "unknown")

                if 'unknown' in type and run.text.strip():
                    type = fix_unknown(run)

                if type == "DefaultParagraphFont":
                    type = fix_DefaultParagraphFont(run)
                    # print paragraph.style.style_id, run.bold, run.font.size, s

                # elif run.bold:          #20.11.16 - Trying to fix 'fake' bold in Appendix
                if run.font.cs_bold:
                    type = bold_type(s, type, run)

                # single run & alignment is CENTER and ...-> letter heading
                if paragraph.alignment is not None and int(paragraph.alignment) == 1 and "heading" not in type:
                    if len(paragraph.runs) <= 2 and run.text.isalnum():
                        size_kind = "heading_letter"
                        type = size_kind

                    if run.font.size and run.text.strip():
                        size_kind = sizes.match(run.font.size)
                        if size_kind == 'unknown':
                            print("!%s. Size: %d, Bool: %s, %s:%s$" % (size_kind, run.font.size, run.font.cs_bold, type, run.text))
                        if size_kind not in ('normal', 'unknown'):
                            type = size_kind

                try:
                    if run.element.rPr.szCs is not None and run.text.strip():
                        type = fix_sz_cs(run, type)

                    if run.element.rPr.bCs is not None and run.text.strip():
                        type = fix_b_cs(run, type)

                    if run.text.strip():
                        type = fix_misc_attrib(run, type)

                    # NOTE: this footnote number need no fix.
                    # it is a recurrence, therefore it has no id.
                    if is_footnote_recurrence(run, type):
                        type = 'footnote_recurrence'
            
                except:
                    pass


                para.append((type, run.text))

                if type == "unknown" and run.text.strip():
                    if run_style_id(run) not in unknown_list:
                        unknown_list.append(run_style_id(run))
                        print(paragraph.text)
                        s = "\nMissing: !%s:%s$\n\n" % (run_style_id(run), run.text)
                        print(s)
                        debug_file.write(s + ' ')


                try:
                    # if run.footnote_references:
                    footnote_references = footnote_run.footnote_references
                    if footnote_references:
                        for (note) in footnote_references:
                            if create_html:
                                html_doc.footnote_ids_of_this_html_doc.append(note.id)
                                relative_note_id = note.id - html_doc.footnote_ids_of_this_html_doc[0] + 1
                                # print "footnote", relative_note_id
                                para.append(('footnote', str(relative_note_id)))
                            elif create_latex:
                                #TODO:  we have a problem here!
                                #what happens in case of both html & latex??
                                para.append(('footnote', str(note.id)))

                except:
                    print("Failed footnote_references")


            para.append(("new_line", "\n"))
            para = analyze_and_fix(para)
            if create_html:
                html_doc = get_active_html_doc(para)
                add_to_output(html_doc, para)
            if create_latex:
                texer.add_to_latex(para, word_doc_footnotes)
        else:
            try:
                # if there is a 'html_doc' - add to id new_line for the paragraph ended
                # if there isn't - it doesn't matter, we're just at the beginning - ignore it
                para = []
                para.append(("new_line", "\n"))
                if create_html:
                    html_doc = html_docs_l[-1]
                    add_to_output(html_doc, para)
                if create_latex:
                    texer.add_to_latex(para, word_doc_footnotes)
            except:
                pass


def add_menu_to_apriory_htmls(html_docs_l):
    # add menus to index.html and search.html
    menu_bar = copy.deepcopy(html_docs_l[0].body.children[0].children[1])
    assert menu_bar['id'] == 'menu_bar'
    content = menu_bar.children[0].children[-1].children[1].children
    del(content[6].attributes['class'])

    place_holder = "<!--menu_bar-->"

    menu_bar_html = menu_bar.render(pretty=False)

    with open("input_web/stub_search.html", 'r', encoding='utf-8') as file:
        menu_bar_html += file.read()

    replace_in_file('output/www/search.html', place_holder, menu_bar_html)

    for index, filename in enumerate((
            'index.html',
            "opening_intros.html",
            "opening_haskamot.html",
            "opening_abbrev.html",
            "opening_signs.html",
    )):
        content[index].attributes['class'] = 'active'
        menu_bar_html = menu_bar.render(pretty=False)
        with open("input_web/stub_search.html", 'r', encoding='utf-8') as file:
            menu_bar_html += file.read()

        replace_in_file('output/www/%s' % filename, place_holder, menu_bar_html)
        del content[index].attributes['class']


if create_html:
    html_docs_l = fix_links(html_docs_l)
    add_menu_to_apriory_htmls(html_docs_l)

    for (html_doc) in html_docs_l:
        close_html_doc(html_doc)

if create_latex:
    texer.close_latex()

with open('output/www/subjects_db.json', 'w', encoding='utf-8') as fp:
    s = json.dumps(subjects_db)
    fp.write("data = " + s)


search_index.create_index()


if unknown_list:
    print("\n\nMissing:")
    print(unknown_list)


def create_zip():
    with zipfile.ZipFile("milon.zip", "w", zipfile.ZIP_DEFLATED) as zf:
        os.chdir("output")
        for dirname, subdirs, files in os.walk("."):
            # avoid creating 'output' directory as first hierrarchy
            # suddenly causes problem with phonegap - makes garbages APKs...
            zf.write(dirname)
            for filename in files:
                if not 'debug' in filename and not '.apk' in filename:
                    zf.write(os.path.join(dirname, filename))
        print("Created milon.zip")
    os.chdir("..")
    shutil.move("milon.zip", "output/")


create_zip()


if process != "ZIP":
    try:
        build_electron.build_electron()
        build_cordova.main()

        if process == "Full":
            playAPISession = upload_google_play.PlayAPISession()
            playAPISession.main(["output/milon.apk"])

    except Exception as e:
        print("Build process failed!")
        print(e)