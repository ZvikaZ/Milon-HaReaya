# hopefully, I will get a new delivery of python-docx supporting szCs
# (see https://github.com/python-openxml/python-docx/issues/248 )
# in the meanwhile, I've hacked it locally
import sys

sys.path.insert(0, r'C:\Zvika\PycharmProjects\python-docx')
sys.path.insert(0, r'C:\Users\sdaudi\Github\python-docx')
import docx  # from aforementioned path

import docx_fork_ludoo
import re
import os

from fixers import fix_sz_cs, fix_b_cs, fix_misc_attrib, fix_unknown, fix_DefaultParagraphFont, fix_section_name
from styles import styles, sizes, bold_type, run_style_id
from helpers import is_subject, is_prev_subject, is_prev_newline, is_prev_meuyan, is_subject_small_or_sub_subject


def is_footnote_recurrence(run, type):
    # a number in superscript, that's not defined as a footnote
    try:
        return \
                run.element.rPr.vertAlign is not None \
                and type != 'footnote' \
                and run.text.strip().isdigit() \
                and list(run.element.rPr.vertAlign.values())[0] == 'superscript'
    except AttributeError:
        return False


def make_sub_subject(subj):
    if subj == 'subject_small':
        return 'sub-subject_normal'
    else:
        return subj


def analyze_and_fix(para):
    # unite splitted adjacent similar types
    prev_type, prev_text = None, ""
    new_para = []
    for (raw_type, text_raw) in para:
        text = text_raw.replace("@", "")
        if text == "◊":
            if prev_text == "":
                # new paragraph with meuayn
                type = "centered_meuyan"
                # see below, end of function, we're going to fix it...
            else:
                # regular meuyan, doesn't need to be centerized
                type = "s02Symbol"
        else:
            type = raw_type
        if prev_type:
            if (type == prev_type) or \
                    (is_subject_small_or_sub_subject(type) and is_subject_small_or_sub_subject(prev_type)) or \
                    (prev_type != "footnote" and prev_type != "footnote_recurrence" and text.strip() in (
                            "", "°", "־", ",")):
                prev_text += text
            elif prev_type == "centered_meuyan" and type == "s02Symbol":
                prev_text += text
            else:
                new_para.append((prev_type, prev_text))
                prev_type, prev_text = type, text
        else:
            prev_type, prev_text = type, text
    new_para.append((prev_type, prev_text))

    # make new_lines stand on their own
    para = new_para
    new_para = []
    for (type, text) in para:
        lines = text.split("\n")
        if len(lines) > 1:
            for (i, line) in enumerate(lines):
                if line:
                    new_para.append((type, line))
                if i + 1 < len(lines):
                    new_para.append(("new_line", "\n"))
        else:
            new_para.append((type, text))

    # fix wrong subjects
    para = new_para
    new_para = []
    for (index, (type, text)) in enumerate(para):
        # if 'subject' in type:
        if is_subject(para, index):
            # real subject is either:
            # first
            # after new_line and empty
            # after subject,"-"
            # after Meuyan
            # second, and after '"'
            if (index == 0) or (is_prev_newline(para, index)) or (is_prev_meuyan(para, index)):
                new_para.append((type, text))
            elif (is_prev_subject(para, index)):
                new_para.append((make_sub_subject(type), text))
            elif new_para[index - 1][0] in ('sub-subject_normal', 'subject_small'):
                new_para.append((make_sub_subject(type), text))
            elif index == 1 and new_para[0][1] == '"':
                # get rid of the wrong '"'
                new_para[0] = ('', '')
                new_para.append((type, '"' + text))
            else:
                new_para.append(("fake_" + type, text))
        elif 'subject' in type:
            # it's got a subject, but 'is_subject' failed
            new_para.append(("fake_" + type, text))
        else:
            new_para.append((type, text))

    # fix missing 'source's
    para = new_para
    new_para = []
    source_pattern = re.compile(r"(.*)(\[.*\])(\..*)")
    for (type, text) in para:
        if source_pattern.match(text) and not 'source' in type:
            g = source_pattern.match(text)
            new_para.append((type, g.group(1)))
            if type in ['definition_small', 'fake_sub-subject_small']:
                new_para.append(('source_normal', g.group(2)))
            elif type == 'definition_normal':
                new_para.append(('source_small', g.group(2)))
            elif 'light' in type:
                new_para.append(('source_light', g.group(2)))
            else:
                print("Fix missing 'source's, unknown type: ", end='')
                print(type, g.group(2))
                new_para.append((type, g.group(2)))
            new_para.append((type, g.group(3)))
        else:
            new_para.append((type, text))

    # make links from circles - °
    para = new_para
    new_para = []
    pattern = re.compile("([\S־]*\S+°)", re.UNICODE)
    for (type, text) in para:
        # don't do this for subjects - it complicates their own link name...
        if "°" in text and 'subject' not in type:
            for (chunk) in pattern.split(text):
                new_para.append((type, chunk))
        else:
            new_para.append((type, text))

    # fix new lines inside headings
    para = new_para
    new_para = []
    ignore_new_line = False
    for (type, text) in para:
        if 'heading' in type:
            ignore_new_line = True
            new_para.append((type, text))
        elif type == "new_line" and ignore_new_line:
            pass
        else:
            new_para.append((type, text))

    # scan for 'empty subjects' ...
    has_subject = False
    has_definition = False
    for (type, text) in para:
        if 'subject' in type and 'fake' not in type and text.strip():
            has_subject = True
            has_definition = False
        elif 'definition' in type:
            has_definition = True

    # ... and fix 'em if required
    if has_subject and not has_definition:
        # empty subject
        # TODO - might be caused by 'heading' interpreted as subject
        para = new_para
        new_para = []
        for (type, text) in para:
            if 'subject' in type and 'fake' not in type:
                new_para.append(('fake_' + type, text))
            else:
                new_para.append((type, text))

    # 'centered_meuyan' is legal only if it's the only "heavy" thing in the paragraph, otherwise, it's a regular 'meuyan"
    found_centered_meuyan = False
    should_replace_centered_meuayn = False
    for (type, text) in para:
        if type == "centered_meuyan":
            found_centered_meuyan = True
        elif type != "new_line" and found_centered_meuyan:
            should_replace_centered_meuayn = True
    if found_centered_meuyan:
        para = new_para
        new_para = []
        if should_replace_centered_meuayn:
            for (type, text) in para:
                if type == "centered_meuyan":
                    new_para.append(("s02Symbol", text))
                else:
                    new_para.append((type, text))
        else:
            # it's a centered meuyan, and should stay such; but we don't need anything else in it
            for (type, text) in para:
                if type == "centered_meuyan":
                    new_para.append((type, text))
                elif type == "new_line":
                    # redundant, and causes troubles...
                    pass
                else:
                    print("Unexpected type in centered meuyan")
                    print(type, text)
                    assert False

    with open('output/debug_fix.txt', 'a', encoding='utf-8') as debug_file:
        debug_file.write("---------------\n")
        for (type, text) in new_para:
            s = "%s:%s.\n" % (type, text)
            debug_file.write(s + ' ')

    # fix
    return new_para


heading_back_to_back = False


# returns:
#  (None, None) - if no need for new section
#  ('NEW_SECTION', string) - with name of new required section
#  ('UPDATE_NAME', string) - to replace the name of newly opened section
#  ('NEW_LETTER', string) - if needs a new section, but w/o putting it in the main TOC
def is_need_new_section(para, prev_name):
    global heading_back_to_back
    for (type, text) in para:
        if type in ("heading_title", "heading_section"):
            if not heading_back_to_back:
                heading_back_to_back = True
                return 'NEW_SECTION', text.strip()
            else:
                # the previous, and this, are headings - unite them
                if prev_name != "מדורים":
                    result = prev_name + " " + text.strip()
                else:
                    # in the special case of 'Section' heading - we don't need it
                    result = text.strip()
                return 'UPDATE_NAME', result
        elif type == "heading_letter":
            try:
                section, _ = prev_name
            except ValueError:
                section = prev_name
            return 'NEW_LETTER', (section, text.strip())

    # if we're here - we didn't 'return text' with a heading
    heading_back_to_back = False
    return None, None


def merge_paras(paras):
    merged_paras = [paras[0]]

    for kind, value in paras[1:]:
        prev_kind, prev_value = merged_paras[-1]

        if prev_kind == kind:
            merged_paras[-1] = (kind, prev_value + value)
        elif value:
            merged_paras.append((kind, value))

    return merged_paras


def parse(doc_file_name):
    def new_page(name, appear_in_toc=True):
        return {
            'name': name,
            'items': [],
            'appear_in_toc': appear_in_toc,
            'footnote_ids': [],
        }

    pages = []
    cur_page = new_page(name=None)

    word_doc = docx.Document(os.path.join('input_dict', doc_file_name))
    word_doc_footnotes = docx_fork_ludoo.Document(os.path.join('input_dict', doc_file_name))

    unknown_list = []

    with open('output/debug.txt', 'w', encoding='utf-8') as debug_file:
        for (paragraph, footnote_paragraph) in zip(word_doc.paragraphs, word_doc_footnotes.paragraphs):
            if paragraph.text.strip():
                # print "Paragraph:", paragraph.text, "$"
                para = []
                debug_file.write("\n\nNEW_PARA:\n------\n")
                for (run, footnote_run) in zip(paragraph.runs, footnote_paragraph.runs):
                    s = "!%s.%s:%s$" % (run.style.style_id, styles.get(run_style_id(run), run_style_id(run)), run.text)
                    # print "!%s:%s$" % (styles.get(run.style.style_id, run.style.style_id), run.text)
                    debug_file.write(s + ' ')
                    type = styles.get(run_style_id(run), "unknown")

                    if 'unknown' in type and run.text.strip():
                        type = fix_unknown(run)

                    if type == "DefaultParagraphFont":
                        type = fix_DefaultParagraphFont(run)
                        # print paragraph.style.style_id, run.bold, run.font.size, s

                    # elif run.bold:          #20.11.16 - Trying to fix 'fake' bold in Appendix
                    if run.font.cs_bold:
                        type = bold_type(s, type, run)

                    # single run & alignment is CENTER and ...-> letter heading
                    if paragraph.alignment is not None and int(paragraph.alignment) == 1 and "heading" not in type:
                        if len(paragraph.runs) <= 2 and run.text.isalnum():
                            size_kind = "heading_letter"
                            type = size_kind

                        if run.font.size and run.text.strip():
                            size_kind = sizes.match(run.font.size)
                            if size_kind == 'unknown':
                                print("!%s. Size: %d, Bool: %s, %s:%s$" % (
                                    size_kind, run.font.size, run.font.cs_bold, type, run.text))
                            if size_kind not in ('normal', 'unknown'):
                                type = size_kind

                    try:
                        if run.element.rPr.szCs is not None and run.text.strip():
                            type = fix_sz_cs(run, type)

                        if run.element.rPr.bCs is not None and run.text.strip():
                            type = fix_b_cs(run, type)

                    except:
                        pass

                    if run.text.strip():
                        type = fix_misc_attrib(run, type)

                    # NOTE: this footnote number need no fix.
                    # it is a recurrence, therefore it has no id.
                    if is_footnote_recurrence(run, type):
                        type = 'footnote_recurrence'

                    para.append((type, run.text))

                    if type == "unknown" and run.text.strip():
                        if run_style_id(run) not in unknown_list:
                            unknown_list.append(run_style_id(run))
                            print(paragraph.text)
                            s = "\nMissing: !%s:%s$\n\n" % (run_style_id(run), run.text)
                            print(s)
                            debug_file.write(s + ' ')

                    try:
                        # if run.footnote_references:
                        footnote_references = footnote_run.footnote_references
                        if footnote_references:
                            for (note) in footnote_references:
                                cur_page['footnote_ids'].append(note.id)
                                relative_note_id = note.id - cur_page['footnote_ids'][0] + 1
                                # TODO texer currently needs 'note.id' and not relative_note_id; keep this, fix in texer
                                para.append(('footnote', str(relative_note_id)))

                    except:
                        print("Failed footnote_references")

                para.append(("new_line", "\n"))
                para = analyze_and_fix(para)

                op, name = is_need_new_section(para, cur_page['name'])
                if op == 'UPDATE_NAME':
                    cur_page['name'] = name
                elif op == 'NEW_LETTER':
                    pages.append(cur_page)
                    cur_page = new_page(name, appear_in_toc=False)
                elif op == 'NEW_SECTION':
                    pages.append(cur_page)
                    cur_page = new_page(fix_section_name(name))
                else:
                    assert op is None and name is None

                cur_page['items'].extend(para)
            else:
                para = [("new_line", "\n")]
                cur_page['items'].extend(para)

    if unknown_list:
        print("\n\nMissing:")
        print(unknown_list)

    for ind, page in enumerate(pages):
        pages[ind]['items'] = merge_paras(page['items'])
        if pages[ind]['items'][0][0] == 'new_line':
            pages[ind]['items'].pop(0)
        try:
            if pages[ind]['items'][-1][0] == 'new_line':
                pages[ind]['items'].pop()
        except IndexError:
            pass

    pages = [page for page in pages if page['items']]
    # drop all pages that have only 1 section of new_line
    pages = [page for page in pages if len(page['items']) > 1 or page['items'][0][0] != 'new_line']

    return pages
